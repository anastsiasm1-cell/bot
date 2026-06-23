"""
Генерация .docx файла с расшифровкой
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt


def build_transcript_docx(text: str, title: str = "Расшифровка") -> io.BytesIO:
    """
    Формирует .docx документ с текстом расшифровки

    Args:
        text: распознанный текст
        title: заголовок документа

    Returns:
        BytesIO с содержимым .docx файла (указатель в начале потока)
    """
    document = Document()
    document.add_heading(title, level=1)

    meta = document.add_paragraph()
    meta_run = meta.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    meta_run.italic = True
    meta_run.font.size = Pt(9)

    document.add_paragraph()

    paragraphs = [p.strip() for p in text.split("\n") if p.strip()] or [text]
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
